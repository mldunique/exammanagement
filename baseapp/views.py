from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login as auth_login, logout
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.db import transaction
from django.utils import timezone
from django.http import JsonResponse
from io import BytesIO
from docx import Document
from random import sample, shuffle
import re, zipfile, os
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from .models import (Subject, Question, Choice, Exam, ExamItem, ExamChoice, 
                    StudentExamSession, StudentAnswer, UserProfile)
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

def login_view(request):
    """Trang đăng nhập chung"""
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        user = authenticate(request, username=username, password=password)
        if user is not None:
            auth_login(request, user)
            # Phân quyền theo role
            try:
                profile = user.userprofile
                if profile.role == 'admin':
                    return redirect('admin_home')
                else:
                    return redirect('student_home')
            except UserProfile.DoesNotExist:
                # Tạo profile mặc định nếu chưa có
                UserProfile.objects.create(user=user, role='student')
                return redirect('student_home')
        else:
            messages.error(request, 'Tên đăng nhập hoặc mật khẩu không đúng')
    
    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

# ===== ADMIN VIEWS =====
# @login_required
def admin_home(request):
    """Trang chủ admin"""
    # Kiểm tra quyền admin
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        return redirect('student_home')
    
    # Lấy thống kê
    subjects = Subject.objects.all()
    all_exams = Exam.objects.select_related('subject').order_by('-created_at')
    recent_exams = all_exams[:5]
    
    # Tính toán thống kê
    stats = {
        'total_subjects': subjects.count(),
        'total_questions': Question.objects.count(),
        'total_exams': all_exams.count(),
        'active_attempts': StudentExamSession.objects.filter(is_submitted=False).count()
    }
    
    return render(request, 'admin_home.html', {
        'subjects': subjects, 
        'exams': recent_exams,
        'all_exams': all_exams,
        'stats': stats
    })

@login_required
@require_http_methods(["GET", "POST"])
def import_docx(request):
    """Import đề thi từ file docx (giữ nguyên code cũ)"""
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        return redirect('student_home')
    
    subjects = Subject.objects.all()
    if request.method == "POST":
        file = request.FILES.get("file")
        if not file:
            messages.error(request, "Hãy chọn file .docx (Template).")
            return redirect("import_docx")

        # Lấy thời gian làm bài từ form
        try:
            duration_minutes = int(request.POST.get("duration_minutes", 60))
            if duration_minutes < 1 or duration_minutes > 300:
                messages.error(request, "Thời gian làm bài phải từ 1 đến 300 phút.")
                return redirect("import_docx")
        except (ValueError, TypeError):
            messages.error(request, "Thời gian làm bài không hợp lệ.")
            return redirect("import_docx")

        raw = file.read()

        # 1) Parse Template.docx (đọc text + image, xử lý nhãn/giá trị xuống dòng)
        try:
            meta, items = _parse_template_docx(raw)
        except Exception as e:
            messages.error(request, "Chỉ chấp nhận file định dạng docx")
            return redirect("import_docx")

        # 2) Map Subject từ header
        subj_raw = (meta['subject'] or '').strip()
        subject = (Subject.objects.filter(code__iexact=subj_raw).first()
                   or Subject.objects.filter(name__iexact=subj_raw).first())
        if not subject:
            messages.error(request, f"Subject '{subj_raw}' không tồn tại. Tạo trước trong admin.")
            return redirect("import_docx")

        # 3) Kiểm tra mã đề và thêm tiền tố subject
        exam_code_raw = (meta['topic_code'] or '').strip()
        if not exam_code_raw:
            messages.error(request, "Thiếu Topic code (Mã đề).")
            return redirect("import_docx")
        
        # Thêm tiền tố subject vào mã đề
        exam_code = f"{subject.code}_{exam_code_raw}"
        if Exam.objects.filter(code=exam_code).exists():
            messages.error(request, f"Mã đề '{exam_code}' đã tồn tại.")
            return redirect("import_docx")

        # Helper: lưu file đúng tên (overwrite nếu trùng tên)
        def save_binary_exact(filename: str, data: bytes) -> str:
            if default_storage.exists(filename):
                default_storage.delete(filename)
            return default_storage.save(filename, ContentFile(data))

        warns = []
        with transaction.atomic():
            created_qs = []

            # 4) Tạo Question/Choice và LƯU ẢNH trực tiếp từ docx theo quy tắc tên
            for q in items:
                qobj = Question.objects.create(
                    subject=subject,
                    text=q.get("text") or "",
                    mark=q.get("mark") or 1.0,
                    unit=q.get("unit") or ""
                )

                # Lưu ảnh (nếu parser có): q["image"] (bytes), q["image_name"] (tên gốc trong docx)
                blob = q.get("image")
                orig = (q.get("image_name") or "").strip()
                if blob:
                    # Lấy phần mở rộng từ tên gốc; fallback .jpg
                    ext = os.path.splitext(orig)[1].lower() or '.jpg'
                    if ext not in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}:
                        ext = '.jpg'

                    # QN từ tài liệu; nếu thiếu, dùng thứ tự hiện tại
                    qn = q.get('id') or (len(created_qs) + 1)
                    # Tên file theo yêu cầu: subjectId_examCode_Q{soThuTu}.{ext}
                    safe_exam = re.sub(r'[^A-Za-z0-9_-]+', '-', exam_code)
                    filename = f"{subject.id}_{safe_exam}_Q{qn}{ext}"

                    saved = save_binary_exact(filename, blob)
                    # ImageField lưu relative path trong MEDIA_ROOT
                    qobj.image.name = saved
                    qobj.save(update_fields=["image"])

                # Lưu lựa chọn A–D
                for label, text in (q.get("choices") or []):
                    Choice.objects.create(
                        question=qobj, label=label, text=text,
                        is_correct=(label == q.get("answer"))
                    )

                created_qs.append((qobj, bool(q.get("mix"))))

            # 5) Tạo Exam + ExamItem + ExamChoice (trộn đáp án nếu mix=True)
            exam = Exam.objects.create(
                code=exam_code,
                subject=subject,
                duration_minutes=duration_minutes,
                question_count=len(created_qs),
            )
            for idx, (qobj, mix) in enumerate(created_qs, start=1):
                item = ExamItem.objects.create(
                    exam=exam, question=qobj, order=idx, mix_choices=mix
                )
                opts = list(qobj.choices.order_by('label'))
                if mix:
                    shuffle(opts)
                labels = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
                for i, opt in enumerate(opts):
                    ExamChoice.objects.create(
                        item=item, label=labels[i], text=opt.text, is_correct=opt.is_correct
                    )

        msg = f"Đã import {len(created_qs)} câu hỏi cho {subject}. Tạo đề '{exam.code}'."
        messages.success(request, msg)
        if warns or '_num_quiz_mismatch' in meta:
            messages.warning(request, "Lưu ý: File có thể không đúng định dạng chuẩn. Vui lòng kiểm tra lại template DOCX.")
        return redirect('exam_preview', exam_id=exam.id)

    return render(request, "import_docx.html", {"subjects": subjects})

@login_required
def exam_create(request):
    """Tạo đề thi ngẫu nhiên"""
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        return redirect('student_home')
    
    subjects = Subject.objects.all()
    if request.method == 'POST':
        code = (request.POST.get('code') or '').strip()
        subject_id = request.POST.get('subject_id')
        duration = int(request.POST.get('duration') or 60)
        n = int(request.POST.get('num_questions') or 10)

        if not code:
            messages.error(request, "Nhập mã đề thi.")
            return redirect('exam_create')

        subject = get_object_or_404(Subject, id=subject_id)
        
        # Thêm tiền tố subject vào mã đề
        exam_code = f"{subject.code}_{code}"
        if Exam.objects.filter(code=exam_code).exists():
            messages.error(request, f"Mã đề '{exam_code}' đã tồn tại.")
            return redirect('exam_create')

        all_qs = list(Question.objects.filter(subject=subject).prefetch_related('choices'))
        if len(all_qs) < n:
            messages.error(request, f"Môn {subject} chỉ có {len(all_qs)} câu, không đủ {n}.")
            return redirect('exam_create')

        picked = sample(all_qs, n)   # chọn ngẫu nhiên n câu
        with transaction.atomic():
            exam = Exam.objects.create(
                code=exam_code, subject=subject,
                duration_minutes=duration, question_count=n
            )
            for idx, q in enumerate(picked, start=1):
                item = ExamItem.objects.create(exam=exam, question=q, order=idx)
                opts = list(q.choices.all())
                shuffle(opts)  # xáo trộn đáp án
                labels = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
                for i, opt in enumerate(opts):
                    ExamChoice.objects.create(
                        item=item,
                        label=labels[i],
                        text=opt.text,
                        is_correct=opt.is_correct  # đúng/sai giữ nguyên theo đáp án gốc
                    )
        return redirect('exam_preview', exam_id=exam.id)

    return render(request, 'exam_create.html', {'subjects': subjects})

@login_required
def exam_preview(request, exam_id):
    """Xem trước đề thi (admin)"""
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        return redirect('student_home')
    
    exam = get_object_or_404(Exam.objects.select_related('subject'), id=exam_id)
    items = exam.items.select_related('question').prefetch_related('choices').order_by('order')
    return render(request, 'exam_preview.html', {'exam': exam, 'items': items})

@login_required
def exam_schedule(request, exam_id):
    """Thiết lập lịch thi (ý 4)"""
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        return redirect('student_home')
    
    exam = get_object_or_404(Exam, id=exam_id)
    
    if request.method == 'POST':
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        is_active = request.POST.get('is_active') == 'on'
        
        if start_time:
            exam.start_time = timezone.datetime.fromisoformat(start_time)
        if end_time:
            exam.end_time = timezone.datetime.fromisoformat(end_time)
        exam.is_active = is_active
        exam.save()
        
        messages.success(request, f"Đã cập nhật lịch thi cho đề {exam.code}")
        return redirect('exam_preview', exam_id=exam.id)
    
    return render(request, 'exam_schedule.html', {'exam': exam})

@login_required
def exam_delete(request, exam_id):
    """Xóa đề thi"""
    if not hasattr(request.user, 'userprofile') or request.user.userprofile.role != 'admin':
        messages.error(request, 'Bạn không có quyền thực hiện thao tác này.')
        return redirect('student_home')
    
    exam = get_object_or_404(Exam, id=exam_id)
    
    if request.method == 'POST':
        exam_code = exam.code
        try:
            # Đếm số lượng dữ liệu liên quan trước khi xóa
            session_count = StudentExamSession.objects.filter(exam=exam).count()
            answer_count = StudentAnswer.objects.filter(session__exam=exam).count()
            item_count = ExamItem.objects.filter(exam=exam).count()
            
            # Xóa đề thi - Django sẽ tự động xóa các bảng liên quan nhờ CASCADE
            exam.delete()
            
            # Thông báo kết quả chi tiết
            if session_count > 0:
                messages.success(request, f"Đã xóa đề thi '{exam_code}' cùng với {session_count} phiên thi, {answer_count} câu trả lời và {item_count} câu hỏi.")
            else:
                messages.success(request, f"Đã xóa đề thi '{exam_code}' cùng với {item_count} câu hỏi.")
            
        except Exception as e:
            messages.error(request, f"Lỗi khi xóa đề thi '{exam_code}': {str(e)}")
        
        return redirect('admin_home')
    
    # GET request - hiển thị trang xác nhận
    session_count = StudentExamSession.objects.filter(exam=exam).count()
    return render(request, 'exam_delete_confirm.html', {
        'exam': exam,
        'session_count': session_count
    })

# ===== STUDENT VIEWS =====
@login_required
def student_home(request):
    """Trang chủ học sinh - danh sách đề thi có thể làm"""
    # Lấy các đề thi có thể làm (đang active và trong thời gian cho phép)
    available_exams = []
    taken_sessions = StudentExamSession.objects.filter(student=request.user).values_list('exam_id', flat=True)
    
    for exam in Exam.objects.filter(is_active=True).order_by('-created_at'):
        if exam.id not in taken_sessions and exam.is_available_now():
            available_exams.append(exam)
    
    # Session đang thi (chưa nộp bài và chưa hết giờ)
    ongoing_sessions = StudentExamSession.objects.filter(
        student=request.user,
        is_submitted=False
    ).select_related('exam', 'exam__subject').order_by('-start_time')
    
    # Lọc session còn thời gian
    active_sessions = []
    for session in ongoing_sessions:
        if not session.is_time_up():
            active_sessions.append(session)
    
    # Lịch sử thi
    completed_sessions = StudentExamSession.objects.filter(
        student=request.user, 
        is_submitted=True
    ).select_related('exam', 'exam__subject').order_by('-end_time')
    
    return render(request, 'student_home.html', {
        'available_exams': available_exams,
        'active_sessions': active_sessions,
        'completed_sessions': completed_sessions
    })

@login_required
def exam_start(request, exam_id):
    """Bắt đầu làm bài thi"""
    exam = get_object_or_404(Exam, id=exam_id)
    
    # Kiểm tra đề có thể làm không
    if not exam.is_available_now():
        messages.error(request, "Đề thi không khả dụng hoặc đã hết hạn")
        return redirect('student_home')
    
    # Kiểm tra đã làm chưa
    if StudentExamSession.objects.filter(student=request.user, exam=exam).exists():
        messages.error(request, "Bạn đã làm đề thi này rồi")
        return redirect('student_home')
    
    # Tạo session mới
    session = StudentExamSession.objects.create(
        student=request.user,
        exam=exam
    )
    
    return redirect('exam_taking', session_id=session.id)

@login_required
def exam_taking(request, session_id):
    """Trang làm bài thi"""
    session = get_object_or_404(StudentExamSession, id=session_id, student=request.user)
    
    # Kiểm tra đã nộp bài chưa
    if session.is_submitted:
        return redirect('exam_result', session_id=session.id)
    
    # Kiểm tra hết giờ chưa - tự động nộp bài
    if session.is_time_up():
        return redirect('exam_submit', session_id=session.id)
    
    # Lấy câu hỏi và câu trả lời hiện tại
    items = session.exam.items.select_related('question').prefetch_related('choices').order_by('order')
    existing_answers = {
        ans.exam_item_id: ans.selected_choice_id 
        for ans in session.answers.select_related('selected_choice')
    }
    
    # Thêm thông tin selected_choice_id vào từng item để template dễ sử dụng
    for item in items:
        item.selected_choice_id = existing_answers.get(item.id)
    
    return render(request, 'exam_taking.html', {
        'session': session,
        'items': items,
        'existing_answers': existing_answers,
        'remaining_time': session.get_remaining_time()
    })

@login_required
def save_answer(request):
    """Lưu câu trả lời (AJAX)"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    session_id = request.POST.get('session_id')
    item_id = request.POST.get('item_id')
    choice_id = request.POST.get('choice_id')
    
    try:
        session = StudentExamSession.objects.get(id=session_id, student=request.user)
        if session.is_submitted or session.is_time_up():
            return JsonResponse({'error': 'Exam is finished'}, status=400)
        
        item = ExamItem.objects.get(id=item_id, exam=session.exam)
        choice = ExamChoice.objects.get(id=choice_id, item=item) if choice_id else None
        
        # Lưu/cập nhật câu trả lời
        answer, created = StudentAnswer.objects.get_or_create(
            session=session,
            exam_item=item,
            defaults={'selected_choice': choice}
        )
        if not created:
            answer.selected_choice = choice
            answer.save()
        
        return JsonResponse({'success': True})
        
    except (StudentExamSession.DoesNotExist, ExamItem.DoesNotExist, ExamChoice.DoesNotExist):
        return JsonResponse({'error': 'Invalid data'}, status=400)

@login_required
def exam_submit(request, session_id):
    """Nộp bài thi"""
    session = get_object_or_404(StudentExamSession, id=session_id, student=request.user)
    
    if session.is_submitted:
        return redirect('exam_result', session_id=session.id)
    
    # Tính điểm
    total_marks = 0
    earned_marks = 0
    
    for item in session.exam.items.select_related('question'):
        total_marks += item.question.mark
        
        try:
            answer = session.answers.get(exam_item=item)
            if answer.selected_choice and answer.selected_choice.is_correct:
                earned_marks += item.question.mark
        except StudentAnswer.DoesNotExist:
            pass  # Câu chưa trả lời = 0 điểm
    
    # Lưu kết quả
    session.end_time = timezone.now()
    session.is_submitted = True
    session.score = earned_marks
    session.total_marks = total_marks
    session.save()
    
    return redirect('exam_result', session_id=session.id)

@login_required
def exam_result(request, session_id):
    """Xem kết quả thi"""
    session = get_object_or_404(StudentExamSession, id=session_id, student=request.user)
    
    if not session.is_submitted:
        return redirect('exam_taking', session_id=session.id)
    
    # Lấy chi tiết câu trả lời
    answers = session.answers.select_related(
        'exam_item__question', 'selected_choice'
    ).order_by('exam_item__order')
    
    results = []
    for answer in answers:
        item = answer.exam_item
        correct_choice = item.choices.filter(is_correct=True).first()
        results.append({
            'question': item.question,
            'selected': answer.selected_choice,
            'correct': correct_choice,
            'is_correct': answer.selected_choice and answer.selected_choice.is_correct,
            'marks': item.question.mark if (answer.selected_choice and answer.selected_choice.is_correct) else 0
        })
    
    return render(request, 'exam_result.html', {
        'session': session,
        'results': results,
        'percentage': round((session.score / session.total_marks) * 100, 1) if session.total_marks > 0 else 0
    })

# ===== HELPER FUNCTIONS (giữ nguyên) =====
def _norm(s: str) -> str:
    if s is None: return ""
    s = s.replace("\xa0", " ")
    s = re.sub(r"[：]", ":", s)  # fullwidth colon -> :
    return re.sub(r"\s+", " ", s).strip()

def _iter_block_items(doc):
    """
    Trả về các block theo đúng thứ tự hiển thị: Paragraph hoặc Table.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def _extract_images_from_paragraph(paragraph):
    """Trả về list [(filename, blob)] các ảnh xuất hiện trong paragraph."""
    out = []
    for run in paragraph.runs:
        # tìm blip (ảnh) trong run
        for blip in run._r.xpath('.//a:blip'):
            rId = blip.get(qn('r:embed'))
            if not rId: 
                continue
            part = paragraph.part.related_parts.get(rId)
            if not part:
                continue
            filename = os.path.basename(part.partname)
            blob = part.blob
            out.append((filename, blob))
    return out

def _doc_to_stream(byte_content):
    """
    Trả về 'dòng sự kiện' theo thứ tự hiển thị:
      {'type':'text', 'text': '...'}
      {'type':'image', 'filename':'xxx.jpg','blob': b'...'}
    Dùng cho cả header và body (bảng/đoạn). Text luôn được đẩy TRƯỚC ảnh để đảm bảo đã có cur.
    """
    doc = Document(BytesIO(byte_content))
    stream = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            # Paragraph: text trước, ảnh sau (đã đúng từ trước)
            txt = _norm(block.text)
            if txt:
                stream.append({'type': 'text', 'text': txt})
            imgs = _extract_images_from_paragraph(block)
            for fn, blob in imgs:
                stream.append({'type': 'image', 'filename': fn, 'blob': blob})
        else:
            # Table: DUYỆT TỪNG HÀNG — TEXT TRƯỚC, ẢNH SAU
            for row in block.rows:
                cells = row.cells

                def cell_text(i):
                    return _norm("\n".join(_norm(p.text) for p in cells[i].paragraphs)) if len(cells) > i else ""

                left_txt  = cell_text(0)
                right_txt = cell_text(1)

                # ---- TEXT PHASE ----
                # QN=... ưu tiên đẩy trước
                if re.match(r'^QN\s*=\s*\d+', left_txt, re.I):
                    stream.append({'type': 'text', 'text': left_txt})
                    if right_txt:
                        stream.append({'type': 'text', 'text': right_txt})
                else:
                    # Lựa chọn a./b./c./d. tách 2 cột
                    if re.match(r'^[A-Da-d][\.\)]$', left_txt) and right_txt:
                        stream.append({'type': 'text', 'text': f"{left_txt} {right_txt}"})
                    elif re.match(r'^[A-Da-d][\.\)]\s+.+', left_txt):
                        stream.append({'type': 'text', 'text': left_txt})
                    # Hàng kiểu ANSWER/MARK/UNIT/MIX chia 2 cột
                    elif re.match(r'^(ANSWER|MARK|UNIT|MIX\s*CHOICES)$', left_txt, re.I) and right_txt:
                        stream.append({'type': 'text', 'text': f"{left_txt}: {right_txt}"})
                    else:
                        # đẩy từng cột nếu có text
                        if left_txt:
                            stream.append({'type': 'text', 'text': left_txt})
                        if right_txt:
                            stream.append({'type': 'text', 'text': right_txt})

                # ---- IMAGE PHASE (sau khi đã đẩy text) ----
                # Ảnh trong từng cell: phải đẩy SAU text để đảm bảo cur đã được tạo (đặc biệt ở hàng QN)
                def push_cell_images(i):
                    if len(cells) > i:
                        for p in cells[i].paragraphs:
                            for fn, blob in _extract_images_from_paragraph(p):
                                stream.append({'type': 'image', 'filename': fn, 'blob': blob})

                push_cell_images(0)
                push_cell_images(1)

    return stream

def _parse_template_docx(byte_content):
    """
    Header: Subject / Number of Quiz / Lecturer / Date / Topic code (chỉ text)
    Body:  QN=<n>, stem (có thể kèm [file:xxx]), options a–d, ANSWER / MARK / UNIT / MIX
           Ảnh: lấy trực tiếp từ .docx; mỗi câu nhận ảnh nhúng đầu tiên gặp sau QN=...
    Trả về:
      meta: dict
      questions: list[{
        id:int, text:str, choices:[(label,text)], answer:'A'..'D',
        image_name:str|None, image:bytes|None, mark:float, unit:str, mix:bool
      }]
    """
    stream = _doc_to_stream(byte_content)  # <- tạo chuỗi sự kiện {'type': 'text'|'image', ...}

    # ---- header ---- (chỉ đọc TEXT đến khi gặp QN=)
    meta = {'subject': None, 'num_quiz': None, 'lecturer': None, 'date': None, 'topic_code': None}
    header_patterns = {
        'subject':    r'^(Subject|Môn\s*học)\s*:\s*(.+)',
        'num_quiz':   r'^(Number\s*of\s*Quiz|Số\s*câu\s*hỏi)\s*:\s*(\d+)',
        'lecturer':   r'^(Lecturer|Giảng\s*viên)\s*:\s*(.+)',
        'date':       r'^(Date|Ngày\s*phát\s*hành)\s*:\s*(.+)',
        'topic_code': r'^(Topic\s*code|Mã\s*đề)\s*:\s*(.+)',
    }
    i = 0
    while i < len(stream):
        ev = stream[i]; i += 1
        if ev['type'] != 'text':
            # header KHÔNG lấy ảnh → bỏ qua
            continue
        ln = _norm(ev['text'])
        if re.match(r'^QN\s*=\s*\d+', ln, re.I):
            i -= 1  # trả lại để vòng sau xử lý như question
            break
        for k, pat in header_patterns.items():
            m = re.match(pat, ln, re.I)
            if m:
                val = m.group(2).strip()
                meta[k] = int(val) if k == 'num_quiz' else val

    missing = [k for k, v in meta.items() if not v]
    if missing:
        raise ValueError("Thiếu thông tin header: " + ", ".join(missing))

    # ---- questions ----
    questions = []
    cur = None
    pending = None  # 'answer' | 'mark' | 'unit' | 'mix'

    while i < len(stream):
        ev = stream[i]; i += 1

        # ẢNH: chỉ gán khi đang ở trong 1 câu hỏi (sau QN=) và chưa có ảnh
        if ev['type'] == 'image':
            if cur and not cur.get("image"):
                cur["image_name"] = ev.get('filename')
                cur["image"] = ev.get('blob')
            continue

        # TEXT
        ln = _norm(ev['text'])
        if not ln:
            continue

        # Bắt đầu câu hỏi
        m_qn = re.match(r'^QN\s*=\s*(\d+)', ln, re.I)
        if m_qn:
            if cur:
                questions.append(cur)
            cur = {
                "id": int(m_qn.group(1)),
                "text": "",
                "choices": [],
                "answer": None,
                "image_name": None,
                "image": None,
                "mark": 1.0,
                "unit": "",
                "mix": False
            }
            pending = None
            continue

        if not cur:
            # chưa vào block QN -> bỏ qua
            continue

        # Nếu đang chờ giá trị cho KEY ở dòng trước (ANSWER/MARK/UNIT/MIX)
        if pending:
            if pending == 'answer' and re.match(r'^[A-D]$', ln, re.I):
                cur['answer'] = ln.upper(); pending = None; continue
            if pending == 'mark' and re.match(r'^[0-9]+(?:\.[0-9]+)?', ln):
                cur['mark'] = float(ln); pending = None; continue
            if pending == 'unit':
                cur['unit'] = ln; pending = None; continue
            if pending == 'mix' and re.match(r'^(Yes|No)$', ln, re.I):
                cur['mix'] = (ln.lower() == 'yes'); pending = None; continue
            # nếu không khớp → rơi xuống như text thường (không consume)

        # Ảnh ghi kiểu [file:xxx] ngay trong text (nếu có) → gán rồi bỏ khỏi text
        img_in_ln = re.findall(r'\[file\s*:\s*([^\]]+)\]', ln, re.I)
        if img_in_ln and not cur.get("image_name"):
            cur["image_name"] = img_in_ln[0].strip()
            ln = re.sub(r'\[file\s*:\s*[^\]]+\]', '', ln, flags=re.I).strip()

        # Lựa chọn A-D
        m_opt = re.match(r'^([A-Da-d])[\.\)]\s*(.+)$', ln)
        if m_opt:
            cur["choices"].append((m_opt.group(1).upper(), m_opt.group(2).strip()))
            continue

        # KEY: value (cùng dòng) HOẶC KEY: (trống) -> bật pending để lấy ở dòng kế tiếp
        m_kv = re.match(r'^(ANSWER|MARK|UNIT|MIX\s*CHOICES)\s*:\s*(.*)$', ln, re.I)
        if m_kv:
            key = m_kv.group(1).upper()
            val = _norm(m_kv.group(2))
            if key.startswith('ANSWER'):
                if val and re.match(r'^[A-D]$', val, re.I): cur['answer'] = val.upper()
                else: pending = 'answer'
            elif key == 'MARK':
                if val: cur['mark'] = float(val)
                else: pending = 'mark'
            elif key == 'UNIT':
                if val: cur['unit'] = val
                else: pending = 'unit'
            else:  # MIX CHOICES
                if val: cur['mix'] = (val.lower() == 'yes')
                else: pending = 'mix'
            continue

        # Thân đề (gộp nhiều dòng)
        cur["text"] = (cur["text"] + ("\n" if cur["text"] else "") + ln).strip()

    if cur:
        questions.append(cur)

    # Đối chiếu số câu
    if meta.get('num_quiz') and meta['num_quiz'] != len(questions):
        meta['_num_quiz_mismatch'] = (meta['num_quiz'], len(questions))

    # Validate từng câu
    for q in questions:
        if len(q["choices"]) < 2:
            raise ValueError(f"Câu QN={q['id']} có ít hơn 2 phương án.")
        if not q.get("answer"):
            raise ValueError(f"Câu QN={q['id']} thiếu ANSWER.")

    return meta, questions

def _peek_next_non_empty(all_lines, i):
    """Trả về (value, index) của dòng kế tiếp KHÔNG rỗng (sau i-1), hoặc (None, i_current)."""
    j = i
    while j < len(all_lines):
        v = _norm(all_lines[j])
        if v:
            return v, j
        j += 1
    return None, i  # không có gì tiếp theo

def _take_if(pattern, text, flags=0):
    m = re.match(pattern, text, flags)
    return m.group(1) if m else None